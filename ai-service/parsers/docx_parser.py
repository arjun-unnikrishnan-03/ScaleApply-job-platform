"""
DOCXParser — extracts text from DOCX resume files.

Uses python-docx for paragraph and table extraction.
Tables are common in modern CV designs (two-column layouts, skill grids)
so we handle them explicitly to avoid losing critical data.
"""

from __future__ import annotations

import logging
from pathlib import Path

from parsers.base import ParsedDocument, ResumeParser
from core.exceptions import EmptyDocumentError, ParseError

logger = logging.getLogger(__name__)

try:
    from docx import Document
    from docx.opc.exceptions import PackageNotFoundError
    _DOCX_AVAILABLE = True
except ImportError:
    _DOCX_AVAILABLE = False


class DOCXParser(ResumeParser):
    """
    Parser for DOCX resume files.

    Extracts text from:
    - Body paragraphs (preserving heading hierarchy)
    - Tables (row-by-row, cell-by-cell)
    - Headers and footers (often contain contact info)
    """

    def supported_extensions(self) -> list[str]:
        return [".docx"]

    def parse(self, file_path: Path) -> ParsedDocument:
        """
        Extract text from DOCX file and return a ParsedDocument.

        Raises:
            ParseError: If the file is not a valid DOCX or cannot be opened.
            EmptyDocumentError: If no text is extractable.
        """
        if not _DOCX_AVAILABLE:
            raise ParseError(
                "python-docx is not installed. Run: pip install python-docx"
            )

        self.validate_path(file_path)

        try:
            doc = Document(str(file_path))
        except PackageNotFoundError as exc:
            raise ParseError(
                message=f"File is not a valid DOCX: {file_path.name}",
                details={"error": str(exc), "path": str(file_path)},
            ) from exc
        except Exception as exc:
            raise ParseError(
                message=f"Failed to open DOCX: {file_path.name}",
                details={"error": str(exc)},
            ) from exc

        sections: list[str] = []

        # ── Body paragraphs ────────────────────────────────────────────────
        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                sections.append(text)

        # ── Tables (two-column layouts, skill matrices, etc.) ─────────────
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(
                    cell.text.strip() for cell in row.cells if cell.text.strip()
                )
                if row_text:
                    sections.append(row_text)

        # ── Headers / Footers (contact info often lives here) ─────────────
        for section in doc.sections:
            for header_footer in [section.header, section.footer]:
                if header_footer:
                    for para in header_footer.paragraphs:
                        text = para.text.strip()
                        if text:
                            sections.append(text)

        raw_text = "\n".join(sections).strip()

        if not raw_text:
            raise EmptyDocumentError(str(file_path))

        logger.debug(
            "DOCX parsed: '%s' | sections=%d | chars=%d",
            file_path.name,
            len(sections),
            len(raw_text),
        )

        return ParsedDocument(
            raw_text=raw_text,
            source_path=str(file_path),
            file_format="docx",
            page_count=len(doc.sections),
        )
